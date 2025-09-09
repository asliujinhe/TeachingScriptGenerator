# app.py
# -*- coding: utf-8 -*-
import io, os, json, zipfile, datetime
from flask import Flask, request, send_file, redirect, url_for, render_template_string, flash
from werkzeug.utils import safe_join

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ROW_HEIGHT_RULE

app = Flask(__name__)
app.secret_key = "change-me"

BASE_DIR = os.path.abspath(os.path.dirname(__file__))
LIB_DIR = os.path.join(BASE_DIR, "jsons")
os.makedirs(LIB_DIR, exist_ok=True)

# 固定 5 个小节标题
FIXED_TITLES = [
    "I.Warming up and Revision",
    "II.Leading-in",
    "Ⅲ. Listening & reading Activities",
    "Ⅳ. Further Development",
    "Ⅴ. Homework",
]

# ---------------- 工具函数 ----------------
def human_size(n: int) -> str:
    if n < 1024: return f"{n} B"
    if n < 1024**2: return f"{n/1024:.1f} KB"
    if n < 1024**3: return f"{n/1024/1024:.1f} MB"
    return f"{n/1024/1024/1024:.1f} GB"

def lib_path(name: str) -> str:
    """在 jsons/ 内安全拼接路径（保留中文/空格），防目录越界。"""
    if not name.lower().endswith(".json"):
        return None
    p = safe_join(LIB_DIR, name)
    if not p:  # None: 越界
        return None
    # 仍防一下越界（极端情况）
    if os.path.commonpath([os.path.abspath(p), LIB_DIR]) != os.path.abspath(LIB_DIR):
        return None
    return p

def next_conflict_name(orig_name: str) -> str:
    """
    名称冲突自动改为 '（n）'（中文全角括号）版。
    例如: 课时.json -> 课时（1）.json -> 课时（2）.json
    """
    base, ext = os.path.splitext(orig_name)
    if ext.lower() != ".json":
        ext = ".json"
    # 去掉末尾已存在的（n）
    import re
    m = re.search(r"（(\d+)）$", base)
    if m:
        base = re.sub(r"（\d+）$", "", base)
    n = 1
    while True:
        candidate = f"{base}（{n}）{ext}"
        if not os.path.exists(os.path.join(LIB_DIR, candidate)):
            return candidate
        n += 1

# ---------------- DOCX 生成（最终版式） ----------------
def enforce_fonts(cell, bold=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(10)
            if bold:
                r.font.bold = True
            rPr = r._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:eastAsia'), 'SimSun')

def align_cell(cell, horiz=None, vert=None):
    if horiz == "center":
        for p in cell.paragraphs:
            p.alignment = 1
    if vert:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), vert)
        tcPr.append(vAlign)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top','bottom','left','right'):
        if edge in kwargs:
            edge_el = tcBorders.find(qn('w:'+edge))
            if edge_el is None:
                edge_el = OxmlElement('w:'+edge)
                tcBorders.append(edge_el)
            for k,v in kwargs[edge].items():
                edge_el.set(qn('w:'+k), v)

def coerce_to_fixed_flow(data):
    src = data.get("教学流程")
    if not isinstance(src, list):
        src = []
    fixed = []
    for i in range(5):
        acts = []
        if i < len(src):
            block = src[i]
            if isinstance(block, dict) and block:
                k0 = list(block.keys())[0]
                v = block.get(k0)
                if isinstance(v, list):
                    for item in v:
                        if isinstance(item, dict):
                            acts.append({"tea": str(item.get("tea","")), "stu": str(item.get("stu",""))})
        fixed.append({FIXED_TITLES[i]: acts})
    return fixed

def json_to_docx_bytes(data: dict, docx_name_hint="lesson_plan") -> bytes:
    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Cm(21)  # A4 宽 21cm
    sec.page_height = Cm(29.7)  # A4 高 29.7cm
    sec.top_margin = Cm(2.54)  # 上 2.54cm
    sec.bottom_margin = Cm(2.54)  # 下 2.54cm
    sec.left_margin = Cm(1.91)  # 左 1.91cm
    sec.right_margin = Cm(1.91)  # 右 1.91cm


    table = doc.add_table(rows=0, cols=3, style="Table Grid")
    table.autofit = True

    # 顶部四项
    for key in ["教学课题", "教学目标", "教学重点与难点", "教学准备"]:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        merged = row_cells[1].merge(row_cells[2])
        merged.text = str(data.get(key, "") or "")
        enforce_fonts(row_cells[0], bold=True); align_cell(row_cells[0], "center", "center")
        enforce_fonts(merged, bold=(key=="教学课题")); align_cell(merged, None, "center")

    # 教学流程标题
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[2]).text = "教 · 学 · 流 · 程"
    enforce_fonts(row_cells[0], bold=True); align_cell(row_cells[0], "center", "center")

    # 表头
    row_cells = table.add_row().cells
    row_cells[0].text = "教师活动"; row_cells[1].text = "学生活动"; row_cells[2].text = "二次备课"
    for c in row_cells: enforce_fonts(c, True); align_cell(c, "center", "center")
    flow_header_row_idx = len(table.rows)-1

    # 教学流程（固定 5 节）
    flow = coerce_to_fixed_flow(data)
    flow_first_content_idx = None
    for idx, block in enumerate(flow):
        title = FIXED_TITLES[idx]
        acts = block.get(title, [])
        row_cells = table.add_row().cells
        p = row_cells[0].paragraphs[0]
        for r in list(p.runs): r.text = ""
        rt = p.add_run(title); rt.bold = True
        p.add_run("\n")
        if acts:
            p.add_run(f"1. {acts[0].get('tea','')}")
            row_cells[1].text = "\n1. " + str(acts[0].get('stu','') or "")
        else:
            row_cells[1].text = ""
        row_cells[2].text = ""
        enforce_fonts(row_cells[0]); enforce_fonts(row_cells[1]); enforce_fonts(row_cells[2])

        if flow_first_content_idx is None:
            flow_first_content_idx = len(table.rows)-1

        for j, step in enumerate(acts[1:], start=2):
            row_cells = table.add_row().cells
            row_cells[0].text = f"{j}. {step.get('tea','')}"
            row_cells[1].text = f"{j}. {step.get('stu','')}"
            row_cells[2].text = ""
            enforce_fonts(row_cells[0]); enforce_fonts(row_cells[1]); enforce_fonts(row_cells[2])

    last_flow_row_idx = len(table.rows)-1

    # 板书设计
    row_cells = table.add_row().cells
    row_cells[0].text = "板书设计"
    merged = row_cells[1].merge(row_cells[2])
    merged.text = str(data.get("板书设计","") or "")
    enforce_fonts(row_cells[0], True); align_cell(row_cells[0], "center", "center")
    enforce_fonts(merged, False)

    # 教学反思（固定空白 3cm）
    row_cells = table.add_row().cells
    row_cells[0].text = "教学反思"
    merged = row_cells[1].merge(row_cells[2]); merged.text = ""
    enforce_fonts(row_cells[0], True); align_cell(row_cells[0], "center", "center")
    table.rows[-1].height = Cm(3); table.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # 边框：仅去掉流程内部横线；竖线/外框保留
    row_count = len(table.rows)
    for r in range(row_count):
        for cell in table.rows[r].cells:
            set_cell_border(cell,
                left={'val':'single','sz':'8','space':'0','color':'000000'},
                right={'val':'single','sz':'8','space':'0','color':'000000'})
    for c in table.rows[flow_header_row_idx].cells:
        set_cell_border(c, bottom={'val':'nil'})
    if flow_first_content_idx is not None and flow_first_content_idx <= last_flow_row_idx:
        for r in range(flow_first_content_idx, last_flow_row_idx):
            for c in table.rows[r].cells: set_cell_border(c, bottom={'val':'nil'})
            for c in table.rows[r+1].cells: set_cell_border(c, top={'val':'nil'})
    for c in table.rows[0].cells: set_cell_border(c, top={'val':'single','sz':'8','space':'0','color':'000000'})
    for c in table.rows[-1].cells: set_cell_border(c, bottom={'val':'single','sz':'8','space':'0','color':'000000'})
    table.alignment = 1

    bio = io.BytesIO()
    doc.save(bio); bio.seek(0)
    return bio.read()

# ---------------- 首页（本地库 + 上传/批量导出） ----------------
INDEX_HTML = """
<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>英语教案编辑器</title>
  <link rel="icon" href="data:,">
  <style>
    body{ font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,"PingFang SC","Hiragino Sans GB","Microsoft YaHei","Helvetica Neue",Arial,sans-serif; margin:2rem auto; max-width:1000px; color:#222; }
    h1{ margin-bottom:.25rem; }
    .muted{ color:#666; }
    .card{ border:1px solid #e5e7eb; border-radius:12px; padding:1rem 1.25rem; margin:1rem 0; box-shadow:0 1px 2px rgba(0,0,0,.04); }
    table{ width:100%; border-collapse:collapse; }
    th,td{ border-bottom:1px solid #eee; padding:.55rem .4rem; font-size:14px; }
    th{ text-align:left; color:#555; }
    .right{ text-align:right; }
    .row{ display:flex; gap:10px; align-items:center; flex-wrap:wrap; }
    .btn{ display:inline-block; border:1px solid #111; padding:.4rem .75rem; border-radius:10px; text-decoration:none; background:#fff; color:#111; cursor:pointer; font-size:14px; }
    .btn:hover{ background:#111; color:#fff; }
    .btn.light{ border-color:#bbb; color:#333; }
    .btn.light:hover{ background:#f5f5f5; color:#111; }
    input[type=file]{ padding:.5rem; border:1px dashed #bbb; border-radius:8px; width:100%; }
    .ok{ color:#076d2d; } .err{ color:#b91c1c; }
  </style>
</head>
<body>
  <h1>英语教案编辑器</h1>
  <p class="muted">管理本地 <code>jsons/</code> 目录里的 JSON格式教案；支持上传、编辑、保存、下载，及勾选批量导出 DOCX/JSON。</p>

  {% with messages = get_flashed_messages(with_categories=true) %}
    {% if messages %}
      {% for cat,msg in messages %}
        <p class="{{ 'ok' if cat=='ok' else 'err' }}">{{ msg }}</p>
      {% endfor %}
    {% endif %}
  {% endwith %}

  <div class="card">
    <h3 style="margin-top:0;">本地 JSON 教案库（jsons/）</h3>
    <form id="bulkExport" action="{{ url_for('export_selected') }}" method="post">
      <table>
        <thead>
          <tr>
            <th style="width:32px;"><input type="checkbox" onclick="toggleAll(this)"></th>
            <th>文件名</th>
            <th>大小</th>
            <th>修改时间</th>
            <th class="right">操作</th>
          </tr>
        </thead>
        <tbody>
          {% for f in files %}
          <tr>
            <td><input type="checkbox" name="selected" value="{{ f.name }}"></td>
            <td>{{ f.name }}</td>
            <td>{{ f.size }}</td>
            <td>{{ f.mtime }}</td>
            <td class="right">
              <a class="btn light" href="{{ url_for('download_json', name=f.name) }}">下载</a>
              <a class="btn" href="{{ url_for('edit_file', name=f.name) }}">编辑</a>
              <a class="btn light" href="{{ url_for('export_one_docx', name=f.name) }}">导出DOCX</a>
            </td>
          </tr>
          {% endfor %}
          {% if not files %}
          <tr><td colspan="5" class="muted">暂无文件</td></tr>
          {% endif %}
        </tbody>
      </table>
      <div class="row" style="justify-content:flex-end; margin-top:.75rem;">
        <button class="btn" name="action" value="docx" type="submit">批量导出 DOCX（ZIP）</button>
        <button class="btn light" name="action" value="json" type="submit">批量下载 JSON（ZIP）</button>
      </div>
    </form>
    <script>
      function toggleAll(cb){
        document.querySelectorAll('input[name=selected]').forEach(x=>x.checked=cb.checked);
      }
    </script>
  </div>

  <div class="card">
    <h3 style="margin-top:0;">上传到本地库（jsons/）</h3>
    <form action="{{ url_for('upload_to_lib') }}" method="post" enctype="multipart/form-data">
      <input type="file" name="files" accept=".json" multiple required>
      <div class="row" style="justify-content:flex-end; margin-top:.75rem;">
        <button class="btn" type="submit">上传到 jsons/</button>
      </div>
    </form>
  </div>
</body>
</html>
"""

# ---------------- 编辑器（沿用你前一版的 flex + 分割线 + 保存/导出按钮） ----------------
EDITOR_HTML = """
<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>英语教案编辑器</title>
  
  <link rel="icon" href="data:,">
  <style>
    :root{ --gap:14px; }
    *{ box-sizing:border-box; }
    body{ font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,"PingFang SC","Hiragino Sans GB","Microsoft YaHei","Helvetica Neue",Arial,sans-serif; margin:1rem auto; max-width:1100px; color:#222; }
    .grid{ display:grid; grid-template-columns:180px 1fr; gap:10px; }
    .card{ border:1px solid #e5e7eb; border-radius:12px; padding:1rem 1.25rem; margin:1rem 0; box-shadow:0 1px 2px rgba(0,0,0,.04); }
    .row{ display:flex; gap:10px; align-items:center; flex-wrap:wrap; }
    .btn{ display:inline-block; border:1px solid #111; padding:.35rem .7rem; border-radius:8px; text-decoration:none; cursor:pointer; font-size:14px; line-height:1.2; }
    .btn.primary{ background:#111; color:#fff; }
    .btn.primary:hover{ background:#222; }
    .btn.secondary{ background:#fff; color:#111; }
    .btn.secondary:hover{ background:#111; color:#fff; }
    .btn.danger{ border-color:#b91c1c; color:#b91c1c; background:#fff; }
    .btn.danger:hover{ background:#b91c1c; color:#fff; }
    .btn.blue{ border-color:#0080FF; color:#0080FF; background:#fff; }
    .btn.blue:hover{ background:#0080FF; color:#fff; }
    input[type=text], textarea{ width:100%; padding:.55rem .65rem; border:1px solid #cfcfcf; border-radius:10px; background:#fff; }
    textarea{ min-height:92px; resize:vertical; }
    #board{ min-height:500px; }
    .muted{ color:#666; }
    .tiny{ font-size:.9rem; color:#777; }
    .pill{ background:#f5f5f5; border-radius:999px; padding:.1rem .6rem; font-size:.85rem; }
    .sec{ border:1px solid #e5e7eb; border-radius:12px; padding:12px; margin:.85rem 0; }
    .sec-hdr{ display:flex; justify-content:space-between; align-items:center; gap:12px; flex-wrap:wrap; }
    .hdr-left{ display:flex; gap:10px; align-items:center; }
    .title{ font-weight:700; }
    .activity-row{ display:flex; align-items:center; gap:var(--gap); padding:0.8rem 0.4rem; flex-wrap:wrap; border-bottom:1px solid #e5e7eb; border-radius:8px; }
    .activity-row:last-child{ border-bottom:none; }
    .activity-row .idx{ flex:0 0 36px; text-align:center; padding-top:.4rem; }
    .activity-row .col{ flex:1 1 360px; min-width:260px; }
    .activity-row .controls{ flex:0 0 160px; display:flex; flex-direction:column; gap:6px; }
    .activity-row .controls .row{ display:flex; gap:8px; }
    .activity-row .controls .row .btn{ flex:1; }
    .activity-row:hover{ background:#f0f0f0; }
    .fab{ position:fixed; right:24px; bottom:24px; background:#0080FF; color:#fff; border:none; border-radius:999px; padding:.9rem 1.2rem; font-size:16px; box-shadow:0 8px 24px rgba(0,0,0,.16); cursor:pointer; z-index:9999; }
    .fab:hover{ background:#2060CF; }
    .save{ position:fixed; right:24px; bottom:82px; background:#10b981; color:#fff; border:none; border-radius:999px; padding:.6rem 1rem; font-size:14px; box-shadow:0 8px 24px rgba(0,0,0,.16); cursor:pointer; z-index:9999; }
    .save:hover{ background:#059669; }
    .back {
  position: fixed;
  right: 24px;
  bottom: 128px; /* 在保存按钮(.save, bottom:82px)上方 */
  background: #f3f4f6;
  color: #111;
  border: 1px solid #d1d5db;
  border-radius: 999px;
  padding: .55rem 1rem;
  font-size: 14px;
  box-shadow: 0 8px 24px rgba(0,0,0,.10);
  cursor: pointer;
  z-index: 9999;
}
.back:hover { background:#e5e7eb; }
  </style>
</head>
<body>
  <h2>英语教案编辑器</h2>
  {% if request.args.get('saved') %}
<div id="save-toast" style="background:#ecfdf5;color:#065f46;border:1px solid #a7f3d0;border-radius:10px;padding:.5rem .75rem;margin:.5rem 0;">
  已保存 ✅
</div>
<script>
  // 2秒后淡出
  setTimeout(()=>{ const t=document.getElementById('save-toast'); if(t){ t.style.transition='opacity .3s'; t.style.opacity='0'; setTimeout(()=>t.remove(),400); } }, 1800);
</script>
{% endif %}
  <p class="muted">小节固定 5 个英文标题（不可编辑）；仅可在各小节内新增/删除/排序“活动”。缺失字段留白。</p>

  <div class="card">
    <div class="grid">
      <label>教学课题</label>
      <input id="kemu" type="text" oninput="state['教学课题']=this.value">
      <label>教学目标</label>
      <textarea id="mubiao" oninput="state['教学目标']=this.value"></textarea>
      <label>教学重点与难点</label>
      <textarea id="zhongdian" oninput="state['教学重点与难点']=this.value"></textarea>
      <label>教学准备</label>
      <input id="zhunbei" type="text" oninput="state['教学准备']=this.value">
    </div>
  </div>

  <div class="card">
    <h3 style="margin:0;">教学流程</h3>
    <div id="secList"></div>
  </div>

  <div class="card">
    <div class="grid">
      <label>板书设计</label>
      <textarea id="board" oninput="state['板书设计']=this.value"></textarea>
    </div>
  </div>

<button class="back" onclick="window.location='{{ url_for('index') }}'">返回主页</button>
  <button class="save" onclick="saveJson()">保存到文件</button>
  <button class="fab" onclick="submitDocx()">生成 DOCX</button>

  <form id="hiddenForm" action="{{ url_for('generate_from_editor') }}" method="post" style="display:none;">
    <input type="hidden" name="json_text" id="json_text">
    <input type="hidden" name="source_filename" id="source_filename" value="{{ filename }}">
  </form>
  <form id="saveForm" action="{{ url_for('save_file') }}" method="post" style="display:none;">
    <input type="hidden" name="json_text" id="save_json_text">
    <input type="hidden" name="source_filename" id="save_source_filename" value="{{ filename }}">
  </form>
  
 


  <script>
    const FIXED = [
      "I.Warming up and Revision",
      "II.Leading-in",
      "Ⅲ. Listening & reading Activities",
      "Ⅳ. Further Development",
      "Ⅴ. Homework"
    ];
    function coerceFlow(data){
      let src = Array.isArray(data['教学流程']) ? data['教学流程'] : [];
      const out = [];
      for (let i=0;i<5;i++){
        let acts=[];
        if(i<src.length && typeof src[i]==='object' && src[i]){
          const k0=Object.keys(src[i])[0]; const arr=src[i][k0];
          if(Array.isArray(arr)){ acts=arr.map(x=>({tea:(x&&x.tea)||"", stu:(x&&x.stu)||""})) }
        }
        const o={}; o[FIXED[i]]=acts; out.push(o);
      }
      return out;
    }

    const state = {{ json_str | safe }};
    state['教学课题']=state['教学课题']||"";
    state['教学目标']=state['教学目标']||"";
    state['教学重点与难点']=state['教学重点与难点']||"";
    state['教学准备']=state['教学准备']||"";
    state['板书设计']=state['板书设计']||"";
    state['教学流程']=coerceFlow(state);

    document.getElementById('kemu').value=state['教学课题'];
    document.getElementById('mubiao').value=state['教学目标'];
    document.getElementById('zhongdian').value=state['教学重点与难点'];
    document.getElementById('zhunbei').value=state['教学准备'];
    document.getElementById('board').value=state['板书设计'];

    function ensureActs(secObj){ const key=Object.keys(secObj)[0]; if(!Array.isArray(secObj[key])) secObj[key]=[]; return key; }
    function addAct(i){ const sec=state['教学流程'][i]; const key=ensureActs(sec); sec[key].push({tea:"",stu:""}); renderSections(); }
    function delAct(i,k){ const sec=state['教学流程'][i]; const key=ensureActs(sec); if(!confirm("确认删除该活动？"))return; sec[key].splice(k,1); renderSections(); }
    function moveAct(i,k,dir){ const sec=state['教学流程'][i]; const key=ensureActs(sec); const j=k+(dir==='up'?-1:1); if(j<0||j>=sec[key].length)return; [sec[key][k],sec[key][j]]=[sec[key][j],sec[key][k]]; renderSections(); }
    function changeTea(i,k,val){ const sec=state['教学流程'][i]; const key=ensureActs(sec); sec[key][k].tea=val; }
    function changeStu(i,k,val){ const sec=state['教学流程'][i]; const key=ensureActs(sec); sec[key][k].stu=val; }

    function renderSections(){
      const box=document.getElementById('secList'); box.innerHTML="";
      state['教学流程'].forEach((secObj,i)=>{
        const key=Object.keys(secObj)[0]; const acts=secObj[key];
        const wrap=document.createElement('div'); wrap.className="sec";
        const hdr=document.createElement('div'); hdr.className="sec-hdr";
        hdr.innerHTML=`<div class="hdr-left"><span class="pill">Section ${i+1}</span><div class="title">${key}</div></div><div class="row"><button class="btn secondary blue" onclick="addAct(${i});return false;">新增活动（Create Activity）</button></div>`;
        wrap.appendChild(hdr);
        acts.forEach((a,k)=>{
          const row=document.createElement('div'); row.className="activity-row";
          row.innerHTML=`
            <div class="idx">${k+1}.</div>
            <div class="col"><div class="tiny">教师活动</div><textarea oninput="changeTea(${i},${k}, this.value)">${(a.tea||"")}</textarea></div>
            <div class="col"><div class="tiny">学生活动</div><textarea oninput="changeStu(${i},${k}, this.value)">${(a.stu||"")}</textarea></div>
            <div class="controls">
              <div class="row">
                <button class="btn secondary" onclick="moveAct(${i},${k},'up');return false;">上移</button>
                <button class="btn secondary" onclick="moveAct(${i},${k},'down');return false;">下移</button>
              </div>
              <button class="btn danger" onclick="delAct(${i},${k});return false;">删除</button>
            </div>`;
          wrap.appendChild(row);
        });
        box.appendChild(wrap);
      });
    }
    renderSections();

    function saveJson(){
      const payload={
        "教学课题":document.getElementById('kemu').value||"",
        "教学目标":document.getElementById('mubiao').value||"",
        "教学重点与难点":document.getElementById('zhongdian').value||"",
        "教学准备":document.getElementById('zhunbei').value||"",
        "教学流程":state['教学流程'],
        "板书设计":document.getElementById('board').value||"",
        "教学反思":""
      };
      document.getElementById('save_json_text').value=JSON.stringify(payload);
      document.getElementById('saveForm').submit();
    }
    function submitDocx(){
      const payload={
        "教学课题":document.getElementById('kemu').value||"",
        "教学目标":document.getElementById('mubiao').value||"",
        "教学重点与难点":document.getElementById('zhongdian').value||"",
        "教学准备":document.getElementById('zhunbei').value||"",
        "教学流程":state['教学流程'],
        "板书设计":document.getElementById('board').value||"",
        "教学反思":""
      };
      document.getElementById('json_text').value=JSON.stringify(payload);
      document.getElementById('hiddenForm').submit();
    }
  </script>
  
</body>
</html>
"""
# ---------------- 路由：主页 ----------------
@app.route("/", methods=["GET"])
def index():
    files = []
    for name in sorted(os.listdir(LIB_DIR)):
        if not name.lower().endswith(".json"): continue
        p = os.path.join(LIB_DIR, name)
        st = os.stat(p)
        files.append({
            "name": name,
            "size": human_size(st.st_size),
            "mtime": datetime.datetime.fromtimestamp(st.st_mtime).strftime("%Y-%m-%d %H:%M"),
        })
    return render_template_string(INDEX_HTML, files=files)

# 上传到库（仅保存到 jsons/，名称冲突自动 “（n）”）
@app.route("/upload_to_lib", methods=["POST"])
def upload_to_lib():
    files = request.files.getlist("files")
    cnt = 0
    for f in files:
        if not f or not f.filename.lower().endswith(".json"):  # 仅 JSON
            continue
        name = os.path.basename(f.filename)  # 保留中文/空格
        path = lib_path(name)
        if not path:
            # 非法名或越界，跳过
            continue
        # 冲突处理
        if os.path.exists(path):
            name = next_conflict_name(name)
            path = lib_path(name)
        f.save(path)
        cnt += 1
    flash(f"已上传 {cnt} 个文件到 jsons/", "ok" if cnt else "err")
    return redirect(url_for("index"))

# 下载单个 JSON
@app.route("/download_json/<path:name>", methods=["GET"])
def download_json(name):
    path = lib_path(name)
    if not path or not os.path.isfile(path):
        flash("文件不存在", "err"); return redirect(url_for("index"))
    return send_file(path, as_attachment=True, download_name=os.path.basename(path), mimetype="application/json")

# 选中项导出（DOCX 或 JSON ZIP）
@app.route("/export_selected", methods=["POST"])
def export_selected():
    selected = request.form.getlist("selected")
    action = request.form.get("action")  # 'docx' or 'json'
    if not selected:
        flash("请至少勾选一个文件", "err")
        return redirect(url_for("index"))

    mem = io.BytesIO()
    with zipfile.ZipFile(mem, "w", zipfile.ZIP_DEFLATED) as zf:
        for name in selected:
            path = lib_path(name)
            if not path or not os.path.isfile(path):
                continue
            if action == "json":
                zf.write(path, arcname=os.path.basename(path))
            else:
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                data["教学流程"] = coerce_to_fixed_flow(data)
                doc_bytes = json_to_docx_bytes(data, docx_name_hint=os.path.splitext(os.path.basename(path))[0])
                zf.writestr(os.path.splitext(os.path.basename(path))[0] + ".docx", doc_bytes)
    mem.seek(0)
    suffix = "docx" if action=="docx" else "json"
    stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    return send_file(mem, as_attachment=True, download_name=f"export_{suffix}_{stamp}.zip", mimetype="application/zip")

# 行内一键导出 DOCX
@app.route("/export_one_docx/<path:name>", methods=["GET"])
def export_one_docx(name):
    path = lib_path(name)
    if not path or not os.path.isfile(path):
        flash("文件不存在", "err"); return redirect(url_for("index"))
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    data["教学流程"] = coerce_to_fixed_flow(data)
    doc_bytes = json_to_docx_bytes(data, docx_name_hint=os.path.splitext(os.path.basename(path))[0])
    return send_file(io.BytesIO(doc_bytes), as_attachment=True,
                     download_name=os.path.splitext(os.path.basename(path))[0] + ".docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# 打开编辑器
@app.route("/edit_file/<path:name>", methods=["GET"])
def edit_file(name):
    path = lib_path(name)
    if not path or not os.path.isfile(path):
        flash("文件不存在", "err"); return redirect(url_for("index"))
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    data = dict(data or {})
    data["教学流程"] = coerce_to_fixed_flow(data)
    # 这里把完整 HTML 送出（你的 EDITOR_HTML 需替换为前面确认的版本）
    return render_template_string(EDITOR_HTML, json_str=json.dumps(data, ensure_ascii=False), filename=os.path.basename(path))

# 保存回库文件
@app.route("/save_file", methods=["POST"])
def save_file():
    text = request.form.get("json_text", "").strip()
    name = request.form.get("source_filename", "").strip()

    # 能回到编辑页就尽量回编辑页；只有文件名都拿不到时才回首页
    def back_to_editor(error_msg=None):
        if error_msg:
            flash(error_msg, "err")
        if name:
            return redirect(url_for("edit_file", name=name, saved=int(error_msg is None)))
        return redirect(url_for("index"))

    path = lib_path(name)
    if not path or not os.path.isfile(path):
        return back_to_editor("保存失败：文件名非法或文件不存在")

    if not text:
        return back_to_editor("保存失败：没有 JSON 内容")

    try:
        data = json.loads(text)
    except Exception as e:
        return back_to_editor(f"保存失败：JSON 解析错误：{e}")

    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    flash(f"已保存到 jsons/{os.path.basename(path)}", "ok")
    return redirect(url_for("edit_file", name=os.path.basename(path), saved=1))


# 从编辑页导出 DOCX
@app.route("/generate_from_editor", methods=["POST"])
def generate_from_editor():
    text = request.form.get("json_text", "").strip()
    name = request.form.get("source_filename", "edited.json").strip() or "edited.json"

    def back_with_error(msg):
        flash(msg, "err")
        # 尽量回到对应的编辑页
        return redirect(url_for("edit_file", name=name)) if name else redirect(url_for("index"))

    if not text:
        return back_with_error("没有收到 JSON 内容")
    try:
        data = json.loads(text)
    except Exception as e:
        return back_with_error(f"JSON 解析失败：{e}")

    doc_bytes = json_to_docx_bytes(data, docx_name_hint=os.path.splitext(name)[0])
    return send_file(
        io.BytesIO(doc_bytes),
        as_attachment=True,
        download_name=os.path.splitext(name)[0] + ".docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    # python app.py
    app.run(host="0.0.0.0", port=5001, debug=True)
